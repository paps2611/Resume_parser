from __future__ import annotations

import io
import re
from typing import Dict, List, Tuple
from docx import Document  # type: ignore

from pydantic import BaseModel

try:
	import docx2txt  # .docx (fallback)
except Exception:  # pragma: no cover
	docx2txt = None  # type: ignore

try:
	import pdfplumber  # .pdf
except Exception:  # pragma: no cover
	pdfplumber = None  # type: ignore


SECTION_KEYWORDS = [
	"summary",
	"objective",
	"experience",
	"employment",
	"work history",
	"education",
	"skills",
	"projects",
	"certifications",
	"awards",
]


class ParsedResume(BaseModel):
	text: str
	sections_present: Dict[str, bool]
	contact_info: Dict[str, str]
	words: List[str]


def _extract_text_from_pdf(file_bytes: bytes) -> str:
	if pdfplumber is None:
		raise RuntimeError("pdfplumber not installed")
	with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
		pages_text = [page.extract_text() or "" for page in pdf.pages]
	return "\n".join(pages_text)


def _extract_text_from_docx(file_bytes: bytes) -> str:
	# Use python-docx for stability; fallback to docx2txt if needed
	try:
		doc = Document(io.BytesIO(file_bytes))
		paras = [p.text for p in doc.paragraphs]
		return "\n".join(paras)
	except Exception:
		if docx2txt is None:
			raise
		try:
			return docx2txt.process(io.BytesIO(file_bytes))  # type: ignore[arg-type]
		except Exception:
			return ""


def extract_text(file_bytes: bytes, filename: str) -> str:
	lower = filename.lower()
	if lower.endswith(".pdf"):
		return _extract_text_from_pdf(file_bytes)
	if lower.endswith(".docx"):
		return _extract_text_from_docx(file_bytes)
	# Plain text fallback
	return file_bytes.decode(errors="ignore")


def parse_resume(file_bytes: bytes, filename: str) -> ParsedResume:
	text = extract_text(file_bytes, filename)
	lower_text = text.lower()
	sections_present = {k: (k in lower_text) for k in SECTION_KEYWORDS}
	words = re.findall(r"[a-zA-Z]+", lower_text)
	contact_info = {}
	# Basic email/phone detection
	email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
	phone_match = re.search(r"(\+\d{1,3}[\s-]?)?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}", text)
	if email_match:
		contact_info["email"] = email_match.group(0)
	if phone_match:
		contact_info["phone"] = phone_match.group(0)
	return ParsedResume(text=text, sections_present=sections_present, contact_info=contact_info, words=words)


def keyword_overlap_score(words: List[str], job_description: str) -> Tuple[int, List[str]]:
	jd_words = re.findall(r"[a-zA-Z]+", job_description.lower())
	if not jd_words:
		return 50, []  # neutral baseline when no JD provided
	jd_set = set(jd_words)
	resume_set = set(words)
	overlap = jd_set & resume_set
	ratio = len(overlap) / max(1, len(jd_set))
	score = int(40 + ratio * 60)  # weight more on overlap
	missing = sorted(list(jd_set - resume_set))[:20]
	return min(score, 100), missing


def formatting_score(text: str) -> int:
	# Penalize images-only or very low text density resumes
	lines = [l for l in text.splitlines() if l.strip()]
	avg_line_len = sum(len(l) for l in lines) / max(1, len(lines))
	if len(lines) < 15:
		return 40
	if avg_line_len < 25:
		return 55
	return 85


def sections_score(sections_present: Dict[str, bool]) -> int:
	present_count = sum(1 for v in sections_present.values() if v)
	return min(100, 50 + present_count * 5)


def contact_score(contact: Dict[str, str]) -> int:
	return 100 if ("email" in contact and "phone" in contact) else 70 if contact else 50


def length_score(text: str) -> int:
	words = len(re.findall(r"\w+", text))
	if words < 200:
		return 60
	if words > 1200:
		return 65
	return 90


def aggregate_score(scores: Dict[str, int]) -> int:
	# Weighted aggregation
	weights = {
		"keyword_match": 0.35,
		"formatting": 0.2,
		"sections": 0.2,
		"contact": 0.1,
		"length": 0.15,
	}
	value = sum(scores[k] * weights[k] for k in weights)
	return int(round(value))


def score_resume_bytes(file_bytes: bytes, filename: str, job_description: str) -> Dict:
	parsed = parse_resume(file_bytes, filename)
	kw_score, missing = keyword_overlap_score(parsed.words, job_description)
	fmt_score = formatting_score(parsed.text)
	sec_score = sections_score(parsed.sections_present)
	con_score = contact_score(parsed.contact_info)
	len_score = length_score(parsed.text)
	breakdown = {
		"keyword_match": kw_score,
		"formatting": fmt_score,
		"sections": sec_score,
		"contact": con_score,
		"length": len_score,
	}
	overall = aggregate_score(breakdown)
	suggestions: List[str] = []
	if kw_score < 75 and missing:
		suggestions.append(
			f"Consider incorporating missing role keywords: {', '.join(missing[:10])}"
		)
	if fmt_score < 70:
		suggestions.append("Increase text density; avoid images/tables; prefer standard headings.")
	if sec_score < 80:
		suggestions.append("Add or improve sections: Summary, Experience, Skills, Education, Projects.")
	if "email" not in parsed.contact_info:
		suggestions.append("Include a professional email address.")
	if "phone" not in parsed.contact_info:
		suggestions.append("Include a reachable phone number with country code.")
	if len_score < 80:
		suggestions.append("Target 1 page (junior) or 1–2 pages (senior) with concise bullets.")
	
	# Enhanced analysis for detailed view
	jd_words = re.findall(r"[a-zA-Z]+", job_description.lower()) if job_description else []
	jd_set = set(jd_words)
	resume_set = set(parsed.words)
	matched_keywords = list(jd_set & resume_set)
	missing_keywords = list(jd_set - resume_set)[:20]
	
	# Section analysis
	section_analysis = {}
	for section, present in parsed.sections_present.items():
		section_analysis[section] = {
			"present": present,
			"score": 100 if present else 0,
			"suggestion": "Good!" if present else f"Consider adding a {section} section"
		}
	
	# Word frequency analysis
	word_freq = {}
	for word in parsed.words:
		if len(word) > 3:  # Only meaningful words
			word_freq[word] = word_freq.get(word, 0) + 1
	
	# Top keywords found
	top_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
	
	# Simplified, always-present fields for UI rendering without feature flags
	return {
		"ats_score": overall,
		"breakdown": breakdown,
		"suggestions": suggestions,
		"extracted": {**parsed.contact_info},
		"resume_text": parsed.text[:4000] + "..." if len(parsed.text) > 4000 else parsed.text,
		"matched_keywords": matched_keywords,
		"missing_keywords": missing_keywords,
		"sections_present": parsed.sections_present,
		"top_keywords": [{"word": word, "count": count} for word, count in top_keywords],
		"word_count": len(parsed.words),
	}


def generate_refined_resume_text(parsed_text: str, job_description: str) -> str:
	from spellchecker import SpellChecker  # type: ignore
	spell = SpellChecker(distance=1)
	jd_words = re.findall(r"[a-zA-Z]+", job_description.lower()) if job_description else []
	core_jd = [w for w in jd_words if len(w) > 3]

	# Normalize whitespace and sentence case
	text = re.sub(r"[\t\r]+", " ", parsed_text)
	lines = [l.strip() for l in text.splitlines() if l.strip()]

	# Spell fix per word where clearly misspelled (keep emails/urls/numbers)
	def fix_spelling(line: str) -> str:
		words = line.split()
		fixed = []
		for w in words:
			if re.search(r"@|https?://|\d", w):
				fixed.append(w)
				continue
			raw = re.sub(r"[^A-Za-z]", "", w)
			if raw and raw.isalpha() and raw.lower() in spell:
				fixed.append(w)
			else:
				candidate = spell.correction(raw.lower()) if raw else None
				fixed.append(w if not candidate or candidate == raw.lower() else w.replace(raw, candidate))
		return " ".join(fixed)

	cleaned = [fix_spelling(l) for l in lines]

	# Ensure key sections
	sections: Dict[str, List[str]] = {
		"summary": [],
		"skills": [],
		"experience": [],
		"education": [],
		"projects": [],
	}
	current = None
	for l in cleaned:
		low = l.lower()
		if any(k in low for k in ["summary", "objective"]):
			current = "summary"; continue
		if "skill" in low:
			current = "skills"; continue
		if any(k in low for k in ["experience", "employment", "work history"]):
			current = "experience"; continue
		if "education" in low:
			current = "education"; continue
		if "project" in low:
			current = "projects"; continue
		if current:
			sections[current].append(l)

	# Auto summary
	if not sections["summary"]:
		sections["summary"] = [f"Experienced professional focusing on {', '.join(core_jd[:6])}."]

	# Skills: aggregate JD keywords and existing lines
	jd_unique = sorted(set(core_jd))
	existing_skills_text = " ".join(sections["skills"]) if sections["skills"] else ""
	merged_skills = sorted(set(re.findall(r"[A-Za-z+#.]+", existing_skills_text.lower()) + jd_unique))
	sections["skills"] = ["• " + ", ".join(merged_skills[:30])]

	# Bulletize experience/projects lines
	def bulletize(items: List[str]) -> List[str]:
		bulleted = []
		for it in items:
			if it.startswith("•"):
				bulleted.append(it)
			else:
				bulleted.append("• " + it.rstrip("."))
		return bulleted
	sections["experience"] = bulletize(sections["experience"]) or ["• Describe your most relevant achievements with metrics."]
	sections["projects"] = bulletize(sections["projects"])

	# Reassemble
	out: List[str] = []
	out += ["Summary"] + sections["summary"] + [""]
	out += ["Skills"] + sections["skills"] + [""]
	if sections["experience"]:
		out += ["Experience"] + sections["experience"] + [""]
	if sections["projects"]:
		out += ["Projects"] + sections["projects"] + [""]
	if sections["education"]:
		out += ["Education"] + sections["education"] + [""]
	return "\n".join(out).strip()


def build_docx_from_text(text: str) -> bytes:
	doc = Document()
	for line in text.splitlines():
		if line.strip() == '':
			doc.add_paragraph('')
			continue
		# Simple headings detection
		low = line.lower().strip()
		if low in SECTION_KEYWORDS or low in ['summary', 'skills']:
			h = doc.add_heading(level=2)
			h.add_run(line.strip())
		else:
			p = doc.add_paragraph(line)
			if line.strip().startswith('•'):
				p.style = doc.styles['List Bullet'] if 'List Bullet' in doc.styles else p.style
	buf = io.BytesIO()
	doc.save(buf)
	return buf.getvalue()




